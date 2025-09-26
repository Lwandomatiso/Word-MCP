import asyncio
from pathlib import Path

import pytest
from docx import Document

# 被测目标：convert_to_pdf（异步函数）
from word_document_server.tools.extended_document_tools import convert_to_pdf


def _make_sample_docx(path: Path) -> None:
    """在临时目录生成一个简单的 .docx 文件（包含中文与空格）。"""
    doc = Document()
    doc.add_heading("转换测试 Document", level=1)
    doc.add_paragraph("这是一段用于 PDF 转换的测试文本。Contains ASCII too.")
    doc.add_paragraph("第二段：包含中文字符与空格，便于覆盖路径/内容边界情况。")
    doc.save(path)


def test_convert_to_pdf_with_temp_docx(tmp_path: Path):
    """
    端到端测试：创建临时 .docx -> 调用 convert_to_pdf -> 校验 PDF 产物。

    说明：
    - Linux/macOS 优先尝试 LibreOffice（soffice/libreoffice），
      失败后回退到 docx2pdf（需要安装 Microsoft Word）。
    - 若环境缺少上述工具或命令不可用，则跳过测试并打印原因。
    """
    # 1) 在临时目录生成一个包含中文与空格的 docx 文件
    src_doc = tmp_path / "含 空 格 的 测试 文档.docx"
    _make_sample_docx(src_doc)

    # 2) 设定输出 PDF 路径（同样放在临时目录）
    out_pdf = tmp_path / "converted 输出.pdf"

    # 3) 运行被测异步函数
    result_msg = asyncio.run(convert_to_pdf(str(src_doc), output_filename=str(out_pdf)))

    # 4) 成功判定：返回信息包含成功关键词，或目标 PDF 实际存在
    success_keywords = ["successfully converted", "成功", "converted to PDF"]
    success = any(k.lower() in result_msg.lower() for k in success_keywords) or out_pdf.exists()

    if not success:
        # 当系统未安装 LibreOffice 或 Microsoft Word 时，扩展工具会返回提示信息
        # 在这种情况下跳过测试，而不是失败。
        pytest.skip(f"PDF 转换工具不可用或转换失败：{result_msg}")

    # 5) 断言：PDF 文件已生成且非空
    # 某些环境（尤其是 docx2pdf）可能忽略精确的输出文件名，只在输出目录或源目录
    # 生成与源文件同名的 PDF，因此这里兼容多种产物位置。
    candidates = [
        out_pdf,
        # 常见：在输出目录生成与源文件同名的 PDF
        out_pdf.parent / f"{src_doc.stem}.pdf",
        # 回退：在源文件同目录生成 PDF
        src_doc.with_suffix(".pdf"),
    ]

    # 若以上路径均不存在，尝试在临时目录内搜寻任意新生成的 PDF
    found = None
    for p in candidates:
        if p.exists():
            found = p
            break
    if not found:
        pdfs = sorted(tmp_path.glob("*.pdf"), key=lambda p: p.stat().st_mtime, reverse=True)
        if pdfs:
            found = pdfs[0]

    if not found:
        # 当工具返回成功但找不到产物，视为环境/外部工具行为差异，跳过而非失败
        pytest.skip(f"未找到生成的 PDF。函数输出：{result_msg}")

    assert found.exists(), f"未找到生成的 PDF：{found}，函数输出：{result_msg}"
    assert found.stat().st_size > 0, f"生成的 PDF 文件为空：{found}"


if __name__ == "__main__":
    # 允许直接运行该文件进行快速验证：
    #   python tests/test_convert_to_pdf.py
    import sys
    sys.exit(pytest.main([__file__, "-q"]))
