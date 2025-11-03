"""
Document creation and manipulation tools for Word Document Server.
"""
import os
import json
import uuid
import httpx
from fastmcp import FastMCP
from io import BytesIO
from typing import Dict, List, Optional, Any
from docx import Document

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension, create_document_copy
from word_document_server.utils.document_utils import get_document_properties, extract_document_text, get_document_structure, get_document_xml, insert_header_near_text, insert_line_or_paragraph_near_text
from word_document_server.core.styles import ensure_heading_style, ensure_table_style


async def create_document(filename: str, title: Optional[str] = None, author: Optional[str] = None) -> str:
    """Create a new Word document with optional metadata.
    
    Args:
        filename: Name of the document to create (with or without .docx extension)
        title: Optional title for the document metadata
        author: Optional author for the document metadata
    """
    filename = ensure_docx_extension(filename)
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot create document: {error_message}"
    
    try:
        doc = Document()
        
        # Set properties if provided
        if title:
            doc.core_properties.title = title
        if author:
            doc.core_properties.author = author
        
        # Ensure necessary styles exist
        ensure_heading_style(doc)
        ensure_table_style(doc)
        
        # Save the document
        doc.save(filename)
        
        return f"Document {filename} created successfully"
    except Exception as e:
        return f"Failed to create document: {str(e)}"


async def get_document_info(filename: str) -> str:
    """Get information about a Word document.
    
    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    try:
        properties = get_document_properties(filename)
        return json.dumps(properties, indent=2)
    except Exception as e:
        return f"Failed to get document info: {str(e)}"


async def get_document_text(filename: str) -> str:
    """Extract all text from a Word document.
    
    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)
    
    return extract_document_text(filename)


async def get_document_outline(filename: str) -> str:
    """Get the structure of a Word document.
    
    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)
    
    structure = get_document_structure(filename)
    return json.dumps(structure, indent=2)


async def list_available_documents(directory: str = ".") -> str:
    """List all .docx files in the specified directory.
    
    Args:
        directory: Directory to search for Word documents
    """
    try:
        if not os.path.exists(directory):
            return f"Directory {directory} does not exist"
        
        docx_files = [f for f in os.listdir(directory) if f.endswith('.docx')]
        
        if not docx_files:
            return f"No Word documents found in {directory}"
        
        result = f"Found {len(docx_files)} Word documents in {directory}:\n"
        for file in docx_files:
            file_path = os.path.join(directory, file)
            size = os.path.getsize(file_path) / 1024  # KB
            result += f"- {file} ({size:.2f} KB)\n"
        
        return result
    except Exception as e:
        return f"Failed to list documents: {str(e)}"


async def copy_document(source_filename: str, destination_filename: Optional[str] = None) -> str:
    """Create a copy of a Word document.
    
    Args:
        source_filename: Path to the source document
        destination_filename: Optional path for the copy. If not provided, a default name will be generated.
    """
    source_filename = ensure_docx_extension(source_filename)
    
    if destination_filename:
        destination_filename = ensure_docx_extension(destination_filename)
    
    success, message, new_path = create_document_copy(source_filename, destination_filename)
    if success:
        return message
    else:
        return f"Failed to copy document: {message}"


async def merge_documents(target_filename: str, source_filenames: List[str], add_page_breaks: bool = True) -> str:
    """Merge multiple Word documents into a single document.
    
    Args:
        target_filename: Path to the target document (will be created or overwritten)
        source_filenames: List of paths to source documents to merge
        add_page_breaks: If True, add page breaks between documents
    """
    from word_document_server.core.tables import copy_table
    
    target_filename = ensure_docx_extension(target_filename)
    
    # Check if target file is writeable
    is_writeable, error_message = check_file_writeable(target_filename)
    if not is_writeable:
        return f"Cannot create target document: {error_message}"
    
    # Validate all source documents exist
    missing_files = []
    for filename in source_filenames:
        doc_filename = ensure_docx_extension(filename)
        if not os.path.exists(doc_filename):
            missing_files.append(doc_filename)
    
    if missing_files:
        return f"Cannot merge documents. The following source files do not exist: {', '.join(missing_files)}"
    
    try:
        # Create a new document for the merged result
        target_doc = Document()
        
        # Process each source document
        for i, filename in enumerate(source_filenames):
            doc_filename = ensure_docx_extension(filename)
            source_doc = Document(doc_filename)
            
            # Add page break between documents (except before the first one)
            if add_page_breaks and i > 0:
                target_doc.add_page_break()
            
            # Copy all paragraphs
            for paragraph in source_doc.paragraphs:
                # Create a new paragraph with the same text and style
                new_paragraph = target_doc.add_paragraph(paragraph.text)
                new_paragraph.style = target_doc.styles['Normal']  # Default style
                
                # Try to match the style if possible
                try:
                    if paragraph.style and paragraph.style.name in target_doc.styles:
                        new_paragraph.style = target_doc.styles[paragraph.style.name]
                except:
                    pass
                
                # Copy run formatting
                for i, run in enumerate(paragraph.runs):
                    if i < len(new_paragraph.runs):
                        new_run = new_paragraph.runs[i]
                        # Copy basic formatting
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        # Font size if specified
                        if run.font.size:
                            new_run.font.size = run.font.size
            
            # Copy all tables
            for table in source_doc.tables:
                copy_table(table, target_doc)
        
        # Save the merged document
        target_doc.save(target_filename)
        return f"Successfully merged {len(source_filenames)} documents into {target_filename}"
    except Exception as e:
        return f"Failed to merge documents: {str(e)}"


async def get_document_xml_tool(filename: str) -> str:
    """Get the raw XML structure of a Word document."""
    return get_document_xml(filename)

# -------------------------------
# 1️⃣ FastAPI app and MCP server
# -------------------------------


# -------------------------------
# 2️⃣ In-memory storage for docs
# -------------------------------
# key: file_id (UUID), value: {"filename": str, "bytes": bytes}
temp_files = {}

# -------------------------------
# 3️⃣ MCP tool: create_document
# -------------------------------
async def create_temp(
    filename: str, 
    title: Optional[str] = None, 
    author: Optional[str] = None
) -> dict:
    """
    Create a Word document in memory and return a temporary download link.

    Returns:
        {"download_url": str, "file_id": str} on success
        {"error": str} on failure
    """
    # Ensure .docx extension
    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    try:
        # Create document
        doc = Document()
        if title:
            doc.core_properties.title = title
        if author:
            doc.core_properties.author = author

        # (Optional) add default styles if needed
        # ensure_heading_style(doc)
        # ensure_table_style(doc)

        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Store in memory with UUID
        file_id = str(uuid.uuid4())
        temp_files[file_id] = {
            "filename": filename,
            "bytes": buffer.getvalue()
        }

        # Return temporary download link (using port 8001 for download server)
        download_url = f"http://127.0.0.1:8001/mcp/download/{file_id}"
        return {"download_url": download_url, "file_id": file_id}

    except Exception as e:
        return {"error": f"Failed to create document: {str(e)}"}

# -------------------------------
# 4️⃣ FastAPI endpoint: download
# -------------------------------
async def load_document_from_url(url: str, filename: str = None) -> dict:
    """
    Load a Word document from a pre-signed URL into memory for editing.
    
    Args:
        url: Pre-signed URL to the .docx file
        filename: Optional custom filename (defaults to the original filename)
        
    Returns:
        {"download_url": str, "file_id": str, "filename": str} on success
        {"error": str} on failure
    """
    import httpx
    
    try:
        # Download the document
        async with httpx.AsyncClient() as client:
            response = await client.get(url)
            response.raise_for_status()
            
            # Check if it's actually a Word document
            content_type = response.headers.get('content-type', '')
            if 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' not in content_type:
                return {"error": "The URL does not point to a valid Word document"}
                
            docx_bytes = response.content
        
        # Generate a filename if not provided
        if not filename:
            content_disp = response.headers.get('content-disposition', '')
            if 'filename=' in content_disp:
                filename = content_disp.split('filename=')[1].strip('"\'')
            else:
                # Try to extract from URL
                from urllib.parse import unquote, urlparse
                path = unquote(urlparse(url).path)
                filename = os.path.basename(path) or f"document_{uuid.uuid4().hex[:8]}.docx"
        
        # Ensure .docx extension
        if not filename.lower().endswith('.docx'):
            filename += '.docx'
            
        # Store in memory
        file_id = str(uuid.uuid4())
        temp_files[file_id] = {
            "filename": filename,
            "bytes": docx_bytes
        }
        
        # Return download URL
        download_url = f"http://127.0.0.1:8001/mcp/download/{file_id}"
        return {
            "download_url": download_url,
            "file_id": file_id,
            "filename": filename
        }
        
    except httpx.HTTPStatusError as e:
        return {"error": f"HTTP error: {str(e)}"}
    except Exception as e:
        return {"error": f"Failed to load document: {str(e)}"}


async def load_example_document() -> dict:
    """
    Load an example Word document from a pre-signed URL into memory.
    
    Returns:
        {"download_url": str, "file_id": str, "filename": str} on success
        {"error": str} on failure
    """
    # Pre-signed URL for the example document
    example_url = "https://flowise-branchai.s3.us-east-1.amazonaws.com/Exampleword.docx?response-content-disposition=inline&X-Amz-Content-Sha256=UNSIGNED-PAYLOAD&X-Amz-Security-Token=IQoJb3JpZ2luX2VjEJv%2F%2F%2F%2F%2F%2F%2F%2F%2F%2FwEaCXVzLWVhc3QtMSJGMEQCIFN4k15gIqv1h7%2F%2FmflgmHRPNWZVBWNnWzV2e%2BG6qrgNAiBOcvXoTwVS0vZRQIblUOIRr3aXQHaYtevbidSzu04CQSq5AwhkEAMaDDgyNzAxMTI0OTgwOCIM52Fo9r7U9IMjfp3vKpYDPvf%2B4LNPunss0d7sJ7ugmT9jnfNK1C7OYFpWlFAzGeiFu7Z%2Ftkd4ZFUmGXCqqb6Q%2F2OXcXpXam33rsOqRF9LJgbU68BirPW1Lad%2FpL4ZLQa8hjLkC2y4nNLKpeSLJ2L7yDgOiVXsTzAHejgdgdQCiQ2i2zIkCHpvRvtZOkekZbhbkJeC5K79Ev7OrNmQmwAiyrJgzwsCr248JAgtsCK4vSbxaMjjyb1ldtbx1PGQYo%2F4ly%2BcSHKkPH%2BVGy9jVm5SIJWv5CHU8etlzXsncl4gVa2D2keykMytD212x1%2FOiElSUuwiL%2Fhy7fgRmsrhyRNSr684wV%2BDaWP%2FevMo6w5SdQP8mi4rNp7cKzUzvlqu4wSgBzSNvwd86OodhoQqNyZF0%2BampPTE2zA6sCOP8HGkL1Zs%2Fps%2FrAH%2Bp21%2BHL3P%2F%2FSnxjYb%2F8y3hRWf5YPlOphKWdE65L0MyCeKnxIHv%2FjaCuX25ZvODU%2FqRa8%2FV1l8Gag0%2FQj%2B8j%2BktT8QPCtOG2nvDg2oGA73A4nWem7e17NUeo8xHN%2FZ%2FDCuh6PIBjrfAuRhvOSUGZkWcIcXslfdO4RzbH0WwAiW13LDaRS5qY4qQg1rQLXK3q4QwOPAP2w2xT1QvkWfTRw2fR%2FL5ysXFOq6my49U68M60GqUH%2BSFH%2BaCejFuRO3Xlkca05tYJVmAhCUdw31CSKhohFu2FSlR%2BusgTbh8GNTAUTwmIVNrSjThTScKR6VmC5sYWFf1hKbraXxGsVnucEtNgApBwXuDAIgG6YRIdrX7MC2bqq5JLzj7OKFcS6omRJPqLH%2FZKdDHR%2Fplp0ZwOVB7KC9f3Fxe8vLCwR7N1mU6TvLRQlJ9NykBCVo8YWSLU5EdxDp16X7ow6PJmq3%2Bc1L5be99FiHmESIZvUSN5se6FDvU0%2BTtMS%2FOqjGrOCXa9gUlatQAO5j1a5w2QisOABgIb5Fr%2FVDP%2B66EqAW7BLUQgqkvmN2nmre3UkF8HwCyotSSPDja20u0y205yEZwcJDthi3pUX6HA%3D%3D&X-Amz-Algorithm=AWS4-HMAC-SHA256&X-Amz-Credential=ASIA4BDNWG2IJFBQENAN%2F20251103%2Fus-east-1%2Fs3%2Faws4_request&X-Amz-Date=20251103T185208Z&X-Amz-Expires=43200&X-Amz-SignedHeaders=host&X-Amz-Signature=975c2656b9e1ba7734d33888d41d48156d39adf63dc2d728864581bd948b6fa2"
    
    try:
        async with httpx.AsyncClient() as client:
            response = await client.get(example_url)
            response.raise_for_status()
            
            # Check if it's actually a Word document
            content_type = response.headers.get('content-type', '')
            if 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' not in content_type:
                return {"error": "The URL does not point to a valid Word document"}
                
            docx_bytes = response.content
        
        # Generate a filename
        filename = "example_document.docx"
            
        # Store in memory
        file_id = str(uuid.uuid4())
        temp_files[file_id] = {
            "filename": filename,
            "bytes": docx_bytes
        }
        
        # Return download URL
        download_url = f"http://127.0.0.1:8001/mcp/download/{file_id}"
        return {
            "download_url": download_url,
            "file_id": file_id,
            "filename": filename
        }
        
    except httpx.HTTPStatusError as e:
        return {"error": f"HTTP error: {str(e)}"}
    except Exception as e:
        return {"error": f"Failed to load example document: {str(e)}"}